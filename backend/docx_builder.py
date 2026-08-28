"""운영자용 요약 보안 진단 보고서(DOCX) 생성기.

원본 전체 결과는 XLSX/JSON으로 제공하고, DOCX는 의사결정에 필요한 취약·검토
항목만 상세히 보여 준다. 따라서 대상/항목 수가 증가해도 불필요하게 수백 페이지로
늘어나지 않는다.
"""
import io
import json
import logging
import os
import re
from collections import Counter, defaultdict

log = logging.getLogger(__name__)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
LOGO_PATH = os.path.join(ASSETS_DIR, "company_logo.png")
SCORE_TABLE_PATH = os.path.join(ASSETS_DIR, "score_calc_table.jpg")
GRADE_TABLE_PATH = os.path.join(ASSETS_DIR, "grade_table.jpg")
ITEM_CATALOG_PATH = os.path.join(ASSETS_DIR, "item_catalog.json")
GUIDE_NAME = "「주요정보통신기반시설 기술적 취약점 분석·평가 방법 상세가이드」(KISA)"
# 보고서 출력 페이지의 "수행 기관" 입력값으로 덮어쓸 수 있는 기본 브랜드명
# (frontend/src/pages/ReportsPage.tsx -> backend/main.py:report()의 org 인자).
ORG_NAME = "HIGHFIVE SECURITY"

NAVY = "0F172A"; BLUE = "2563EB"; SLATE = "475569"; LIGHT = "F8FAFC"; BORDER = "CBD5E1"
HEADER_FILL = "E2E8F0"; HEADER_TEXT = (30, 41, 59)  # 표 헤더행: 진하게 칠하지 않고 연한 회색 + 어두운 글자로
CMP_LABELS = {"fixed": "조치완료", "still_vuln": "여전히 취약", "new": "신규 발견", "regressed": "악화"}


def status_key(value):
    s = str(value or "").strip().upper()
    if s in {"양호", "OK", "GOOD", "PASS"}: return "양호"
    if s in {"취약", "FAIL", "VULNERABLE"}: return "취약"
    if s in {"N/A", "NA", "NOT APPLICABLE", "해당없음"}: return "N/A"
    return "검토"  # 수동확인, 예외, 판정불가 포함


def effective_status(result: dict) -> str:
    """manual_verdict(관리자가 '검토' 항목을 양호/취약으로 확정한 값)가 있으면
    그걸 우선한다 - db.py의 recompute_host_score(점수 계산)가 이미 이 규칙을
    쓰는데, 보고서 통계/표는 원래 status만 보고 있어서 확정한 뒤에도 계속
    '검토'로 잡히는 불일치가 있었다. status 원본 컬럼 자체는 DB에서 안
    건드린다(자동 진단 결과 기록 보존) - 보고서에 보여줄 값만 여기서 맞춘다."""
    return status_key(result.get("manual_verdict") or result.get("status"))


def clean(value, default="-"):
    """DB 텍스트/HTML의 공백을 정리해 표에 그대로 노출되는 깨짐을 줄인다."""
    s = str(value or "").replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" *\n *", "\n", s).strip()
    return s or default


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{name}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa"); tcMar.append(el)
    tcPr.append(tcMar)


def set_borders(table):
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="{BORDER}"/>'
        f'<w:left w:val="single" w:sz="4" w:color="{BORDER}"/><w:bottom w:val="single" w:sz="4" w:color="{BORDER}"/>'
        f'<w:right w:val="single" w:sz="4" w:color="{BORDER}"/><w:insideH w:val="single" w:sz="4" w:color="{BORDER}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="{BORDER}"/></w:tblBorders>'))


def font(run, size=9, bold=False, color=(30, 41, 59)):
    run.font.name = "Malgun Gothic"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor(*color)


def kill_list(p):
    """이 문단에는 글머리 기호/번호 매기기를 절대 적용하지 않는다고 명시적으로
    박아 둔다(w:numId=0). LibreOffice에서는 멀쩡히 표 셀/일반 문단으로
    보이는데 실제 MS Word(맑은 고딕)로 열면 표의 모든 행·일부 문단 앞에
    검은 사각형 글머리 기호(■)가 붙어 나오는 문제가 실측됨 - 문서 XML
    어디에도 numPr을 명시한 적이 없는데도 Word가 리스트로 표시한 것이라,
    스타일 상속 쪽에서 원인을 특정하기보다 문단마다 "리스트 없음"을 직접
    선언해 Word의 판단 여지를 없앤다."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)


def write_cell(cell, value, size=9, bold=False, color=(30, 41, 59), align=None):
    cell.text = clean(value, "")
    margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
    kill_list(p)
    for run in p.runs: font(run, size, bold, color)


def no_split(table, keep_together=False):
    """행 내부가 페이지 경계에서 잘리는 건 항상 막는다(cantSplit).

    keep_together=True면 표 전체를 한 덩어리로 취급해서(마지막 행 제외 모든
    행에 keep_with_next) 페이지에 다 안 들어가면 표째로 다음 페이지로
    넘긴다 - 카드/요약처럼 몇 행 안 되는 짧은 표에서 "표 중간에 페이지가
    끊기는 것"을 막을 때만 쓴다. 수십~수백 행짜리 표(점검 항목 카탈로그,
    호스트 목록 등)에 걸면 반대로 표 전체가 안 들어갈 때마다 앞 페이지에
    큰 빈 공백을 남기고 통째로 다음 페이지로 밀려버리는 부작용이 있어서
    (실측됨) 기본값은 False이고, 그런 긴 표는 repeat_header로 자연스럽게
    여러 페이지에 걸치게 두는 쪽이 맞다."""
    rows = table.rows
    n = len(rows)
    for i, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr(); trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if keep_together and i < n - 1:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True


def paragraph(doc, text="", size=9.5, bold=False, color=(51, 65, 85), align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = align; p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before); p.paragraph_format.line_spacing = 1.2
    kill_list(p)
    r = p.add_run(clean(text, "")); font(r, size, bold, color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12); p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    kill_list(p)
    r = p.add_run(text); font(r, 16 if level == 1 else 12, True, (15, 23, 42))
    return p


def _border_bottom(p, color=BORDER, sz=6):
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/></w:pBdr>'))


def add_letterhead(doc, org=ORG_NAME):
    """표지(1p)는 비워두고, 그 뒤 본문 페이지에만 로고 + 대외비 표기가 있는
    머리말을 반복 표시한다."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    content_width = section.page_width - section.left_margin - section.right_margin

    p = header.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_after = Pt(4)
    kill_list(p)
    if os.path.exists(LOGO_PATH):
        p.add_run().add_picture(LOGO_PATH, width=Inches(0.28))
    r = p.add_run(f"  {org}"); font(r, 9, True, (30, 41, 59))
    r2 = p.add_run("\t대외비 · CONFIDENTIAL"); font(r2, 8, True, (153, 27, 27))
    _border_bottom(p)


def add_page_footer(doc, org=ORG_NAME):
    section = doc.sections[0]; footer = section.footer; p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kill_list(p)
    r = p.add_run(f"{org}  |  "); font(r, 8, False, (100, 116, 139))
    for typ, text in (("begin", None), ("instr", "PAGE"), ("end", None)):
        el = OxmlElement("w:fldChar" if typ != "instr" else "w:instrText")
        if typ != "instr": el.set(qn("w:fldCharType"), typ)
        else: el.set(qn("xml:space"), "preserve"); el.text = text
        r._r.append(el)
    paragraph(footer, f"본 보고서는 {GUIDE_NAME} 기준을 준용하여 작성되었습니다. 이 문서는 대외비이며, 수신자 외 제3자에게 무단으로 공유·배포할 수 없습니다.",
              7.5, False, (148, 163, 184), WD_ALIGN_PARAGRAPH.CENTER, after=0, before=2)


TOC_TITLE = "목  차"


def add_toc(doc):
    # 여기서는 자리만 잡아둔다(제목 문단 + 빈 문단). 실제 목차는
    # docx_toc.refresh_fields()가 헤드리스 LibreOffice로 이 문서를 열어 그
    # 빈 문단 자리에 진짜 목차 인덱스 객체를 만들고 채운 뒤 다시 저장한다 -
    # python-docx가 만든 TOC 필드(fldSimple/fldChar)는 LibreOffice가 자기
    # 내부의 "목차 인덱스" 객체로 인식을 못 해 갱신이 안 되는 걸 확인했다.
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kill_list(p)
    r = p.add_run(TOC_TITLE); font(r, 18, True, (15, 23, 42))
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(15)
    kill_list(p)


def data_summary(hosts):
    totals = Counter(); risk = Counter(); categories = defaultdict(Counter); repeated = defaultdict(list); priority = []
    for h in hosts:
        for r in h.get("results", []):
            st = effective_status(r); imp = clean(r.get("importance"), "중")
            totals[st] += 1; categories[clean(r.get("category"), "기타")][st] += 1
            if st == "취약":
                risk[imp] += 1; repeated[(clean(r.get("code")), clean(r.get("title")))].append(h.get("hostname", "-"))
                priority.append((0 if imp == "상" else 1 if imp == "중" else 2, h, r))
    priority.sort(key=lambda x: x[0])
    return totals, risk, categories, repeated, priority


def action_breakdown(hosts):
    """'취약/검토'로 잡힌 건들이 조치 관점에서 어느 단계에 있는지 집계한다.

    이 보고서의 모든 수치는 스캔 시점의 최종 상태(자동 조치 스크립트가
    실행됐다면 그 재검증 결과까지 반영된 값) 기준이다 - 최초 진단 때 취약이던
    항목이 스캔 도중 조치돼 양호로 바뀌었으면 그 즉시 진단 결과 자체가
    양호로 덮어써지기 때문에(apply_remediation_result), "조치 전/후" 두
    시점을 이 표 하나로 동시에 보여줄 수는 없다. 대신 reviewed 플래그로
    "한 번도 손 안 댐(미조치)"과 "조치를 시도했지만 안 고쳐짐(재조치 필요)"을
    구분해서, 지금 남아있는 취약 건수 중 실제로 손을 댔는데도 안 고쳐진
    게 얼마나 되는지는 드러낸다. 스캔 간(이전 회차 대비) 진짜 전/후 비교는
    3.5. 조치 경과 섹션이 담당한다."""
    counts = Counter()
    for h in hosts:
        for r in h.get("results", []):
            st = effective_status(r)
            if r.get("fixed_by_user"):
                counts["조치 완료"] += 1
            elif r.get("manual_verdict"):
                counts["수동 판정"] += 1
            elif st == "취약":
                counts["재조치 필요" if r.get("reviewed") else "미조치"] += 1
            elif st == "검토":
                counts["검토 필요"] += 1
    return counts


def auto_summary(project, hosts, totals, risk, categories, scan):
    """AI 종합 소견(consultant_comment)이 없을 때 쓰는 규칙 기반 총평. ANTHROPIC_API_KEY
    미설정/생성 실패 시 consultant_comment가 빈 문자열로 남는 경우가 실제로 흔해서
    (02_generate_report.py 참고), 총평 문단 자체가 통째로 빠지지 않도록 항상 실제
    집계값(data_summary 결과, DB 캐시 컬럼 아님 - audit_scans.total_checks 등은
    실측과 어긋나는 경우가 있어 신뢰하지 않는다)으로 문장을 만들어 대체한다."""
    total_checks = sum(totals.values())
    parts = [
        f"이번 진단은 {project}를 대상으로 {len(hosts)}대 서버, 총 {total_checks}개 점검 항목을 기준으로 수행되었습니다.",
        f"종합 보안 점수는 {clean(scan.get('average_security_score'))}점({clean(scan.get('total_grade'))} 등급)이며, "
        f"취약 {totals['취약']}건(이 중 위험도 '상' {risk['상']}건), 검토 필요 {totals['검토']}건이 확인되었습니다.",
    ]
    top_cat, top_stat = max(categories.items(), key=lambda kv: kv[1]["취약"], default=(None, None))
    if top_cat and top_stat["취약"] > 0:
        parts.append(f"특히 '{top_cat}' 영역에서 취약점이 {top_stat['취약']}건으로 가장 많이 발견되어 우선적인 점검이 필요합니다.")
    if hosts:
        worst = min(hosts, key=lambda h: float(h.get("security_score_100") or 0))
        parts.append(f"보안 점수가 가장 낮은 서버는 {clean(worst.get('hostname'))}({clean(worst.get('security_score_100'))}점)로, 우선 조치 대상 서버로 권고합니다.")
    return " ".join(parts)


def summary_table(doc, totals, risk):
    """양호/취약/검토 등을 색으로 구분하지 않고, 문서 다른 표들과 같은 톤(남색
    헤더 + 무채색 본문)의 평범한 표로 보여준다."""
    table = doc.add_table(rows=2, cols=7); table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_borders(table)
    labels = ["양호", "취약", "검토", "N/A", "상 위험", "중 위험", "하 위험"]
    counts = [totals["양호"], totals["취약"], totals["검토"], totals["N/A"], risk["상"], risk["중"], risk["하"]]
    for i, label in enumerate(labels):
        write_cell(table.cell(0, i), label, 8.5, True, HEADER_TEXT, WD_ALIGN_PARAGRAPH.CENTER); shade(table.cell(0, i), HEADER_FILL)
    for i, count in enumerate(counts):
        write_cell(table.cell(1, i), f"{count}건", 14, True, (30, 41, 59), WD_ALIGN_PARAGRAPH.CENTER); shade(table.cell(1, i), LIGHT)
    no_split(table, keep_together=True)


def simple_table(doc, headers, rows, widths=None, keep_together=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_borders(table)
    for i, h in enumerate(headers):
        write_cell(table.cell(0, i), h, 8.5, True, HEADER_TEXT, WD_ALIGN_PARAGRAPH.CENTER); shade(table.cell(0, i), HEADER_FILL)
        if widths: table.cell(0, i).width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            write_cell(cells[i], val, 8.5, False, (51, 65, 85), WD_ALIGN_PARAGRAPH.CENTER if i in (0, len(row) - 1) else None)
            if widths: cells[i].width = Inches(widths[i])
    # keep_together 지정이 없으면 행 수로 자동 판단 - 몇 줄 안 되는 표는
    # 페이지 중간에서 쪼개지지 않게 붙이고, 긴 표(카탈로그·호스트 목록 등)는
    # repeat_header와 함께 자연스럽게 여러 페이지에 걸치도록 둔다.
    if keep_together is None:
        keep_together = len(rows) <= 10
    no_split(table, keep_together=keep_together)
    return table


def add_figure(doc, path, width, caption=None):
    """참고 자료 이미지를 그대로 삽입한다(표로 다시 그리지 않고 원본 그림 사용).
    파일이 없으면(자산 미배포 등) 조용히 건너뛴다 - 이미지 하나 때문에 보고서
    생성 전체가 실패하면 안 된다."""
    if not os.path.exists(path):
        log.warning("참고 이미지가 없어 건너뜁니다: %s", path)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kill_list(p)
    p.add_run().add_picture(path, width=Inches(width))
    if caption:
        paragraph(doc, caption, 8, False, (100, 116, 139), WD_ALIGN_PARAGRAPH.CENTER, before=2, after=10)


PIPELINE_STEPS = [
    ("①", "스크립트 배포", "SSH로 전달"),
    ("②", "자동 점검", "KISA 항목별 실행"),
    ("③", "결과 수집", "JSON → DB 저장"),
    ("④", "자동 조치", "가능한 항목만"),
    ("⑤", "재검증", "결과 갱신"),
    ("⑥", "점수·등급 산출", "배점 기준(2장)"),
    ("⑦", "보고서 생성", "XLSX, DOCX, JSON"),
]


def add_pipeline_diagram(doc):
    """진단이 실제로 어떻게 수행되는지(자동화 스크립트 파이프라인)를 표 기반
    흐름도로 보여준다. 이미지 라이브러리 없이 표 셀 음영/테두리만으로 그려서
    - matplotlib/Pillow 같은 새 런타임 의존성이 필요 없다."""
    n = len(PIPELINE_STEPS)
    ncols = n * 2 - 1  # 단계 칸 + 화살표 칸 번갈아
    step_w, arrow_w = 0.83, 0.2  # 7*step_w + 6*arrow_w ≈ 7.01in, 본문 폭(7.06in) 안에 맞춤
    table = doc.add_table(rows=2, cols=ncols); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (num, label, desc) in enumerate(PIPELINE_STEPS):
        col = i * 2
        top = table.cell(0, col); bot = table.cell(1, col)
        top.merge(bot)
        write_cell(top, f"{num} {label}\n{desc}", 8, False, (15, 23, 42), WD_ALIGN_PARAGRAPH.CENTER)
        top.paragraphs[0].runs[0].bold = True
        margins(top, 70, 70, 40, 40); shade(top, LIGHT)
        tcPr = top._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="{BORDER}"/>'
                               f'<w:left w:val="single" w:sz="4" w:color="{BORDER}"/><w:bottom w:val="single" w:sz="4" w:color="{BORDER}"/>'
                               f'<w:right w:val="single" w:sz="4" w:color="{BORDER}"/></w:tcBorders>'))
        table.columns[col].width = Inches(step_w)
        if i < n - 1:
            acol = col + 1
            atop = table.cell(0, acol); abot = table.cell(1, acol); atop.merge(abot)
            write_cell(atop, "→", 11, True, (100, 116, 139), WD_ALIGN_PARAGRAPH.CENTER)
            table.columns[acol].width = Inches(arrow_w)
    no_split(table, keep_together=True)
    return table


def repeat_header(table, nrows=1):
    """긴 표가 페이지를 넘어갈 때 헤더 행이 각 페이지 상단에 다시 나오게 한다."""
    for row in table.rows[:nrows]:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))


def load_item_catalog():
    """items.sh(U-01~U-67, D-01~D-26) 메타데이터 + KISA 가이드 PDF에서 뽑은 짧은
    설명을 담은 정적 참고 자료. 진단 결과(호스트 수)와 무관하게 82행 고정이라,
    대상 서버가 수천 대로 늘어나도 이 표는 그대로다 - 서버별 상세 대신 이 표로
    "무엇을 점검하는지"의 레퍼런스를 제공한다."""
    if not os.path.exists(ITEM_CATALOG_PATH):
        log.warning("점검 항목 카탈로그가 없어 건너뜁니다: %s", ITEM_CATALOG_PATH)
        return []
    with open(ITEM_CATALOG_PATH, encoding="utf-8") as f:
        return json.load(f)


def add_comparison_section(doc, comparisons):
    """3.5. 조치 경과 - 호스트별로 기준 회차(baseline) 대비
    이번 회차의 조치완료/여전히 취약/신규 발견/악화 현황을 요약한다.
    comparisons는 backend/db.py::get_comparison_data()의 반환값 그대로."""
    keys = ["fixed", "still_vuln", "new", "regressed"]
    for h in comparisons:
        title = f"{h['hostname']} ({h['ip']})"
        if h["is_baseline"]:
            paragraph(doc, f"{title} - 최초 진단 회차라 비교 대상 없음", 9, False, (100, 116, 139), after=8)
            continue

        # 기준↔현재 회차 사이 실제 경과일 - "조치 경과"라는 섹션명에 맞게
        # 며칠 만에 이 변화가 있었는지 같이 보여준다. 둘 다 날짜가 있을 때만
        # 계산한다(host_facts에 scan_date가 없을 극히 예외적인 경우 대비).
        before_d, after_d = h.get("before_scan_date"), h.get("after_scan_date")
        days_str = f" · {(after_d - before_d).days}일 경과" if before_d and after_d else ""
        paragraph(doc, f"{title}  —  기준({h['before_scan_id']}) → 현재({h['after_scan_id']}){days_str}", 9.5, True, (30, 41, 59), after=4)

        table = doc.add_table(rows=2, cols=4); table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_borders(table)
        for i, key in enumerate(keys):
            write_cell(table.cell(0, i), CMP_LABELS[key], 9, True, HEADER_TEXT, WD_ALIGN_PARAGRAPH.CENTER); shade(table.cell(0, i), HEADER_FILL)
            write_cell(table.cell(1, i), f"{len(h[key])}건", 12, True, (30, 41, 59), WD_ALIGN_PARAGRAPH.CENTER); shade(table.cell(1, i), LIGHT)
        no_split(table, keep_together=True)

        # 코드 목록은 너무 길어지지 않게 최대 10개 + 나머지는 건수만 표기
        for key in keys:
            items = h[key]
            if not items:
                continue
            codes_str = ", ".join(f"{it['code']}({it['title']})" for it in items[:10])
            if len(items) > 10:
                codes_str += f" 외 {len(items) - 10}건"
            paragraph(doc, f"· {CMP_LABELS[key]}: {codes_str}", 8.5, False, (71, 85, 105), after=4)

        empty = doc.add_paragraph(); empty.paragraph_format.space_after = Pt(6); kill_list(empty)


def generate_docx(full_data, comparisons=None, org=None, inspector=None, customer=None):
    org = clean(org, ORG_NAME)
    doc = Document(); section = doc.sections[0]
    section.top_margin = Inches(.75); section.bottom_margin = Inches(.7); section.left_margin = Inches(.72); section.right_margin = Inches(.72)
    styles = doc.styles; styles["Normal"].font.name = "Malgun Gothic"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    add_page_footer(doc, org)
    add_letterhead(doc, org)
    scan = full_data.get("scan") or {}
    # DB 조회 순서(등록/스캔 순서)가 아니라 hostname 사전순으로 통일한다 -
    # XLSX(csv_builder.generate_xlsx)는 이미 이렇게 정렬해서 시트를 만드는데
    # DOCX는 원래 순서 그대로라 두 산출물의 서버 순서가 어긋날 수 있었다.
    hosts = sorted(full_data.get("hosts") or [], key=lambda h: h.get("hostname") or "")
    totals, risk, categories, repeated, priority = data_summary(hosts)
    project = clean(scan.get("project_name"), "주요정보통신기반시설 시스템 취약점 진단")

    # Cover
    # [MOD] 표지 하단의 "대외비/KISA 준용" 안내 박스(notice_table)가 Word에서
    # 실제로 열어보면(맑은 고딕 폰트 메트릭 - LibreOffice 렌더링보다 줄 높이가
    # 커짐) 1페이지에 다 안 들어가고 2페이지로 밀려나는 게 실측됐다. 로고
    # 위 여백과 프로젝트명 아래 여백을 줄여 1페이지 안에 들어오게 한다.
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph(); p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(50)
        kill_list(p_logo)
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(0.95))
    paragraph(doc, org, 13, True, (37, 99, 235), WD_ALIGN_PARAGRAPH.CENTER, after=18, before=(14 if os.path.exists(LOGO_PATH) else 75))
    paragraph(doc, "보안 취약점 진단\n결과 보고서", 27, True, (15, 23, 42), WD_ALIGN_PARAGRAPH.CENTER, after=14)
    paragraph(doc, project, 12, False, (100, 116, 139), WD_ALIGN_PARAGRAPH.CENTER, after=40)
    # [MOD] "회사명"(진단 대상 고객사) - "수행 기관"(org, 진단하는 쪽)만 있고
    # 진단받는 회사명이 어디에도 없어서 어느 회사 보고서인지 모호하다는 피드백으로
    # 추가함. 표 맨 위(가장 먼저 보이는 행)에 둔다.
    cover_rows = [["회사명", clean(customer)], ["스캔 ID", clean(scan.get("scan_id"))], ["진단 일시", clean(scan.get("scan_date"))], ["종합 점수 / 등급", f"{clean(scan.get('average_security_score'))}점 / {clean(scan.get('total_grade'))}"], ["점검 대상", f"총 {len(hosts)}대 서버"]]
    # 보고서 출력 페이지에서 채워 넣은 경우에만 표시 - "1.2. 수행 인력" 팀
    # 명단(auditor)과는 별개로, 이 보고서를 요청/발주받은 담당자 1명을 표지에
    # 추가로 밝히고 싶을 때 쓰는 값이라 팀 표를 덮어쓰지 않고 별도 행으로 둔다.
    if inspector:
        cover_rows.append(["진단 담당자", clean(inspector)])
    cover = simple_table(doc, ["구분", "내용"], cover_rows, [1.5, 4.8])
    paragraph(doc, f"핵심 결과  |  취약 {totals['취약']}건 · 검토 {totals['검토']}건 · 상 위험 {risk['상']}건", 11, True, (153, 27, 27), WD_ALIGN_PARAGRAPH.CENTER, before=16)

    # [MOD] AI 종합 소견(audit_scans.consultant_comment)은 표지에 따로 안 보여준다 -
    # "3.1. 총평"이 같은 값을 그대로 쓰는데(없으면 auto_summary로 대체), 표지에도
    # 전문을 그대로 실으면 완전히 같은 문단이 표지·본문에 두 번 나오는 중복이었다
    # (실측 확인됨). 표지는 위 "핵심 결과" 통계 한 줄로 충분하고, 상세 서술은
    # 본문 3.1절 한 곳에만 남긴다.

    notice_table = doc.add_table(rows=1, cols=1); notice_table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_borders(notice_table)
    notice_cell = notice_table.cell(0, 0); notice_cell.width = Inches(5.5)
    write_cell(notice_cell, f"본 보고서는 {GUIDE_NAME} 기준을 준용하여 작성되었습니다.\n"
               "본 문서에 기재된 내용은 대외비(CONFIDENTIAL)로 분류되며, 발주기관 및 관련 담당자 외 제3자에게 "
               "사전 서면 동의 없이 공유·복제·배포할 수 없습니다.", 8, False, (100, 116, 139), WD_ALIGN_PARAGRAPH.CENTER)
    notice_cell.paragraphs[0].paragraph_format.line_spacing = 1.35
    shade(notice_cell, "FAFAFA"); no_split(notice_table)

    doc.add_page_break()

    add_toc(doc); doc.add_page_break()

    heading(doc, "1. 진단 개요", 1)
    paragraph(doc, f"본 보고서는 {project}의 일환으로 수행된 서버·데이터베이스 기술적 취약점 진단 결과를 정리한 문서입니다. "
              "주요정보통신기반시설 보안 수준을 객관적으로 점검하고, 발견된 취약점에 구체적인 조치 방안을 제시하여 침해사고를 예방하고 "
              "정보보호 수준을 향상시키는 데 목적이 있습니다. 이어지는 장에서는 점수 산출 기준, 종합 진단 현황, 우선 조치 대상을 순서대로 "
              "다루며, 항목별 전체 원본 결과는 별도 XLSX 파일에서 코드로 조회할 수 있습니다.", after=10)

    heading(doc, "1.1. 진단 절차 및 방법론", 2)
    paragraph(doc, "KISA 가이드 항목에 매핑된 자동화 스크립트로 서버 OS·DBMS 설정을 점검하고, "
              "자동 조치가 가능한 항목은 조치 스크립트 실행 후 재검증까지 수행합니다. 스크립트로 자동 판정이 어려운 항목은 "
              "'검토'로 표시되며, 관리자가 현재 설정(점검 증적)을 직접 확인해 최종 판정(수동 판정)합니다.", after=10)
    add_pipeline_diagram(doc)
    paragraph(doc, "④ 자동 조치는 스크립트가 있는 항목에서만 실행되며, ⑤ 재검증 이후에도 취약 상태가 남으면 "
              "XLSX 전체 결과에 '재조치 필요'로 별도 표시됩니다.", 8, False, (100, 116, 139), before=6, after=10)

    heading(doc, "1.2. 수행 인력", 2)
    auditor_raw = clean(scan.get("auditor"), "")
    auditors = [a.strip() for a in re.split(r"[,/;]", auditor_raw) if a.strip()] if auditor_raw and auditor_raw != "-" else []
    # 1.1절 진단 절차(점검 → 조치 → 재검증 → 보고서)에 맞춰 역할을 나눈다.
    # 목록에 없는 이름(팀 구성이 바뀐 경우)은 "진단 수행"으로 기본 배정한다.
    DUTY_BY_NAME = {
        "심수용": "프로젝트 총괄",
        "김성진": "진단 수행",
        "김하영": "진단 수행",
        "정진우": "조치 및 재검증",
        "한주협": "조치 및 재검증",
    }
    if auditors:
        rows = [[org, name, DUTY_BY_NAME.get(name, "진단 수행")] for name in auditors]
        simple_table(doc, ["소속", "성명", "담당 업무"], rows, [1.8, 1.6, 1.6])
    else:
        paragraph(doc, "등록된 수행 인력 정보가 없습니다.", 9, False, (100, 116, 139))

    doc.add_page_break()
    heading(doc, "2. 점수 산출 기준", 1)
    heading(doc, "2.1. 배점 및 판정 기준", 2)
    paragraph(doc, "점검 항목은 중요도에 따라 배점이 다르며, 진단 결과에 따라 다음과 같이 점수를 적용합니다.", after=6)
    add_figure(doc, SCORE_TABLE_PATH, 5.2, f"[{GUIDE_NAME} 붙임4 발췌]")
    paragraph(doc, "본 시스템의 판정은 양호·취약·검토·N/A 네 가지이며(부분조치 'P' 등급은 별도로 두지 않음), "
              "위 표의 O(취약 발견)/×(취약 제거) 두 행에 해당하는 배점만 사용합니다.", 8.5, False, (100, 116, 139), after=10)
    heading(doc, "2.2. 산출식", 2)
    paragraph(doc, "호스트별 보안 점수는 아래 산출식으로 계산합니다(위 가이드의 '기술적 취약점 점수 계산식'과 동일).", after=6)
    formula_table = doc.add_table(rows=1, cols=1); formula_table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_borders(formula_table)
    fcell = formula_table.cell(0, 0); fcell.width = Inches(5.2)
    write_cell(fcell, "보안 점수 = (배점 합계 − 취약 항목 배점 합계) ÷ 배점 합계 × 100", 11, True, (15, 23, 42), WD_ALIGN_PARAGRAPH.CENTER)
    shade(fcell, "F1F5F9"); no_split(formula_table)
    paragraph(doc, "'검토' 항목은 관리자가 수동 판정을 확정하면 그 판정값(양호/취약)을 우선 적용하고, 미확정 상태에서는 배점에는 포함하되 감점하지 않습니다. "
              "N/A 항목은 배점 대상에서 제외합니다. 본 시스템은 서버·DBMS 기술적 진단 전용이므로, 위 가이드의 관리적·물리적 영역 합산 및 "
              "망분리 비율 가중치는 적용하지 않습니다(기술적 취약점 점수를 그대로 최종 점수로 사용).", 8.5, False, (100, 116, 139), before=8, after=10)
    heading(doc, "2.3. 등급 기준", 2)
    paragraph(doc, "산출된 보안 점수(100점 만점 → 비율)에 따라 아래 기준으로 등급을 부여합니다.", after=6)
    # 2.1의 배점표(SCORE_TABLE_PATH)와 마찬가지로 가이드 원문에서 그대로 발췌한
    # 표라 이미지로 유지한다 - 직접 다시 그린 표로 바꾸면 가이드 원문 그대로임을
    # 보장할 수 없다. 원 안의 등급별 색(●) 자체는 이미지에 포함된 내용이라
    # 코드에서 건드릴 수 없다.
    add_figure(doc, GRADE_TABLE_PATH, 4.4, f"[{GUIDE_NAME} 등급 기준표]")

    catalog = load_item_catalog()
    if catalog:
        # 위 가이드 자체가 "Unix 서버"와 "Ⅷ. DBMS"를 별개 장으로 다루므로, 그
        # 구성을 그대로 따라 표도 U-계열/D-계열을 나눈다 - 대상 서버 수와
        # 무관하게 각각 항목 수가 고정돼 있는 정적 레퍼런스라는 점은 동일하다.
        unix_items = [c for c in catalog if c["code"].startswith("U-")]
        db_items = [c for c in catalog if c["code"].startswith("D-")]
        intro = (f"본 진단에 사용된 전체 점검 항목입니다. 대상 서버 수와 무관하게 항목 수는 고정되어 있으며, "
                 "호스트별 실제 진단 결과(현재 설정, 취약/양호 여부, 조치 권고)는 별도 XLSX 전체 결과에서 코드로 조회할 수 있습니다. "
                 f"설명은 {GUIDE_NAME}을 요약한 것입니다.")
        if unix_items:
            heading(doc, f"2.4. UNIX 점검 항목 (U-01~U-{len(unix_items):02d})", 2)
            paragraph(doc, intro, 9, False, (100, 116, 139), after=8)
            rows = [[c["code"], c["category"], c["title"], c["description"]] for c in unix_items]
            table = simple_table(doc, ["코드", "점검 영역", "항목명", "설명"], rows, [.55, .95, 1.7, 3.85])
            repeat_header(table)
        if db_items:
            heading(doc, "2.5. DBMS 점검 항목 (D-01~D-26)", 2)
            if not unix_items:
                paragraph(doc, intro, 9, False, (100, 116, 139), after=8)
            rows = [[c["code"], c["category"], c["title"], c["description"]] for c in db_items]
            table = simple_table(doc, ["코드", "점검 영역", "항목명", "설명"], rows, [.55, .95, 1.7, 3.85])
            repeat_header(table)

    doc.add_page_break()
    heading(doc, "3. 종합 진단 현황", 1)
    heading(doc, "3.1. 총평", 2)
    # AI 종합 소견(consultant_comment)이 있으면 그대로, 없으면(API 키 미설정/생성
    # 실패로 실제로 빈 경우가 흔함) 실측 집계로 만든 총평으로 대체 - 총평 문단
    # 자체가 통째로 빠지는 일이 없게 한다.
    summary_text = (scan.get("consultant_comment") or "").strip() or auto_summary(project, hosts, totals, risk, categories, scan)
    paragraph(doc, summary_text, 9.5, False, (51, 65, 85), after=10)
    heading(doc, "3.2. 종합 진단 통계", 2); summary_table(doc, totals, risk)
    action = action_breakdown(hosts)
    action_labels = ["미조치", "재조치 필요", "조치 완료", "수동 판정", "검토 필요"]
    if any(action.get(k) for k in action_labels):
        paragraph(doc, "위 수치는 조치 스크립트 실행 및 재검증까지 반영된 최종 상태 기준입니다. "
                  "'재조치 필요'는 조치를 시도했음에도 여전히 취약한 항목으로, 손도 안 댄 '미조치'와는 구분됩니다.",
                  8.5, False, (100, 116, 139), before=8, after=6)
        rows = [[k, action.get(k, 0)] for k in action_labels if action.get(k)]
        action_table = simple_table(doc, ["조치 상태", "건수"], rows, [2.0, 1.2])
    heading(doc, "3.3. 호스트별 현황", 2)
    host_rows = []
    for i, h in enumerate(hosts, 1):
        rs = h.get("results", []); vuln = sum(effective_status(r) == "취약" for r in rs); review = sum(effective_status(r) == "검토" for r in rs)
        host_rows.append([i, clean(h.get("hostname")), clean(h.get("ip")), clean(h.get("os")), f"{clean(h.get('security_score_100'))}점", f"취약 {vuln} / 검토 {review}"])
    host_table = simple_table(doc, ["No.", "호스트명", "IP 주소", "OS", "보안 점수", "조치 대상"], host_rows, [.45, 1.05, 1.15, 1.25, .8, 1.25])
    repeat_header(host_table)  # 대상 서버가 많아져 표가 여러 페이지로 넘어가도 헤더가 반복 표시됨
    heading(doc, "3.4. 점검 영역별 취약 현황", 2)
    category_rows = []
    for cat, cnt in sorted(categories.items(), key=lambda x: x[1]["취약"], reverse=True):
        category_rows.append([cat, cnt["취약"], cnt["검토"], cnt["양호"], cnt["N/A"]])
    simple_table(doc, ["점검 영역", "취약", "검토", "양호", "N/A"], category_rows, [2.8, .8, .8, .8, .8])

    if comparisons is not None:
        heading(doc, "3.5. 조치 경과", 2)
        add_comparison_section(doc, comparisons)

    doc.add_page_break()
    heading(doc, "4. 우선 조치 권고", 1)
    paragraph(doc, "위험도 '상' 취약점과 여러 서버에서 반복적으로 발견된 취약점을 우선 조치 대상으로 제시합니다.", after=8)
    top_rows = []
    for _, h, r in priority[:15]: top_rows.append([clean(r.get("importance"), "중"), clean(r.get("code")), clean(r.get("title")), f"{clean(h.get('hostname'))} ({clean(h.get('ip'))})"])
    if top_rows: simple_table(doc, ["위험도", "코드", "취약점", "대상 서버"], top_rows, [.75, .75, 2.8, 2.0])
    else: paragraph(doc, "현재 '취약'으로 판정된 항목이 없습니다.", 9, False, (100, 116, 139))
    heading(doc, "4.1. 반복 취약점", 2)
    repeated_rows = []
    for (code, title), names in sorted(repeated.items(), key=lambda x: len(x[1]), reverse=True):
        if len(names) < 2:
            continue
        # 서버 수가 많아지면 이름 나열이 무한히 길어지므로 XLSX 요약 표(반복
        # 취약점 Top5)와 같은 기준으로 최대 10개 + 나머지는 건수만 표기한다.
        names_str = ", ".join(names[:10]) + (f" 외 {len(names) - 10}대" if len(names) > 10 else "")
        repeated_rows.append([code, title, names_str, f"{len(names)}대"])
    if repeated_rows: simple_table(doc, ["코드", "취약점", "발견 서버", "영향 범위"], repeated_rows, [.75, 2.5, 2.2, .75])
    else: paragraph(doc, "2대 이상에서 반복된 취약점이 없습니다.", 9, False, (100, 116, 139))

    output = io.BytesIO(); doc.save(output); return output.getvalue()
