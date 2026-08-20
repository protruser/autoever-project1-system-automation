import io
from docx import Document

def generate_docx(full_data):
    doc = Document()
    scan = full_data.get("scan", {})
    hosts = full_data.get("hosts", [])

    doc.add_heading(scan.get("project_name", "KISA 보안 진단 보고서"), level=0)
    doc.add_paragraph(f"스캔 ID: {scan.get('scan_id')} | 진단 일시: {scan.get('scan_date')}")
    doc.add_paragraph(f"종합 점수: {scan.get('average_security_score')}점 ({scan.get('total_grade')})")

    for h in hosts:
        doc.add_heading(f"호스트: {h.get('hostname')} ({h.get('ip')})", level=1)
        doc.add_paragraph(f"OS: {h.get('os')} | 점수: {h.get('security_score_100')}점 | 준수율: {h.get('compliance_rate')}")
        
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "코드", "항목명", "상태", "중요도"

        for r in h.get("results", []):
            row = table.add_row().cells
            row[0].text = r.get("code", "")
            row[1].text = r.get("title", "")
            row[2].text = r.get("status", "")
            row[3].text = r.get("importance", "")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
